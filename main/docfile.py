from itertools import product
import os
import re
from typing import Optional

import docx # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.shared import Pt # type: ignore
from docx.text.paragraph import Paragraph # type: ignore

from main.utils import convert_dates, get_locale, paragraph_replace_text


FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(11)
HEADER_TOP_MARGIN = Pt(24)
HEADER_BTM_MARGIN = Pt(6)


class DocFile(object):
    def __init__(
        self, 
        input_path: str, 
        output_path: str, 
        file: str
    ) -> None:
        self.input_path = input_path
        self.output_path = output_path
        self.file = file
        self.doc = docx.Document(os.path.join(self.input_path, file))

    def replace_text(
        self, 
        old_string: str, 
        new_str: str, 
        language: str,
        date_format: Optional[str] = None
    ) -> None:
        MONTHS, new_quotes = get_locale(language)
        regex = re.compile(re.escape(old_string))

        def replace_quotes(paragraph: Paragraph) -> Paragraph:
            paragraph.text = re.sub(r'\"(.*?)\"', new_quotes, paragraph.text)
            return paragraph
        
        for paragraph in self.doc.paragraphs:
            paragraph_replace_text(paragraph, regex, new_str)
            replace_quotes(paragraph)
            if date_format:
                paragraph.text = convert_dates(paragraph.text, MONTHS, date_format)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_replace_text(paragraph, regex, new_str)
                        replace_quotes(paragraph)
                        if date_format:
                            paragraph.text = convert_dates(paragraph.text, MONTHS, date_format)

        for section in self.doc.sections:
            for p1, p2, p3 in product(
                section.header.paragraphs, 
                section.first_page_header.paragraphs,
                section.footer.paragraphs
            ):
                paragraph_replace_text(p1, regex, new_str)
                paragraph_replace_text(p2, regex, new_str)
                paragraph_replace_text(p3, regex, new_str)

                replace_quotes(p1)
                replace_quotes(p2)
                replace_quotes(p3)

                if date_format:
                    p1.text = convert_dates(p1.text, MONTHS, date_format)
                    p2.text = convert_dates(p2.text, MONTHS, date_format)
                    p3.text = convert_dates(p3.text, MONTHS, date_format)

    def add_start_text(self, start_text: Optional[str] = None) -> None:
        if start_text:
            section = self.doc.sections[0]
            section.different_first_page_header_footer = True
            existing_text = '\n'.join(p.text for p in section.header.paragraphs)
            first_page_header = section.first_page_header
            first_paragraph = first_page_header.paragraphs[0]
            run = first_paragraph.add_run(start_text)
            first_page_header.add_paragraph()
            first_page_header.paragraphs[-1].add_run(existing_text)

            paragraph_format = first_paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            section.header_distance = HEADER_TOP_MARGIN
            paragraph_format.space_after = HEADER_BTM_MARGIN
            
            font = run.font
            font.name = FONT_NAME
            font.size = FONT_SIZE
            font.italic = True
            font.underline = True

    def save_file(self) -> None:
        self.doc.save(os.path.join(self.output_path, self.file))