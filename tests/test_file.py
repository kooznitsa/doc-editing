import docx
import io
import os
from pathlib import Path
import pytest
import textwrap
import uuid

from werkzeug.datastructures import FileStorage # type: ignore

import app
from main.config import ALLOWED_EXTENSIONS
from main.file import File


class MockFile(File):
    def __init__(self, input_path, output_path):
        super().__init__(uuid.uuid4())
        self.input_path = input_path
        self.output_path = output_path
        self.archive_name = 'test_archive.zip'


class TestFile:
    ORIG_CONTENT = textwrap.dedent(
        """
        Tesla's diary  contains just one comment
        ( on what happened at the end of his employment )
        - a note he scrawled across the two pages covering
        7 December 1884 to 4 January 1885 , saying 
        "Good by to the Edison Machine Works".
        """
    )
    EDITED_CONTENT = textwrap.dedent(
        """
        Tesla's diary contains just one comment
        (on what happened at the end of his employment)
        — a note he scrawled across the two pages covering
        1884-12-07 to 1885-01-04, saying 
        “Good by to the Edison Machine Works”.
        """
    )
    
    filenames = ('foo.docx', 'bar.docx', 'baz.docx')
    wrong_filenames = ('wrong1.txt', 'wrong2.png')

    @pytest.fixture(scope='session')
    def create_directories(self, tmp_path_factory: Path):
        input_directory = tmp_path_factory.mktemp('tmp')
        output_directory = tmp_path_factory.mktemp('tmp')

        for filename in self.filenames:
            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.text = self.ORIG_CONTENT
            document.save(os.path.join(input_directory, filename))

        yield input_directory, output_directory

    @pytest.fixture(scope='module')
    def test_client(self):
        flask_app = app.app
        testing_client = flask_app.test_client()
        ctx = flask_app.app_context()
        ctx.push()
        yield testing_client
        ctx.pop()

    def test_allowed_file(self):
        for filename in self.wrong_filenames:
            assert '.' in filename and filename.rsplit('.', 1)[1].lower() not in ALLOWED_EXTENSIONS
    
    def test_upload(self, test_client):
        data = FileStorage(
            stream=io.BytesIO(b'some initial text data'),
            filename='fake-text-stream.docx',
        )
        response = test_client.post('/upload', data=data)
        assert response.status_code == 200

    def test_download(self, test_client):
        data = FileStorage(
            stream=io.BytesIO(b'some initial text data'),
            filename='fake-text-stream.docx',
        )
        response = test_client.post('/download', data=data)
        assert response.status_code == 200

    def get_text(self, path):
        doc = docx.Document(path)
        return '\n'.join(para.text for para in doc.paragraphs)

    def test_replace_text(self, create_directories):
        input_directory, output_directory = create_directories

        for filename in os.listdir(input_directory):
            mock_file = MockFile(input_directory, output_directory)
            mock_file.edit_file(filename, 'English', '%Y-%m-%d', '---First line---')
            file = output_directory / filename

            assert self.get_text(file) == self.EDITED_CONTENT